"""
Pariana Backend — Contracts Router (Fully Refactored)

Handles:
  - POST /api/upload              → Upload and process a PDF contract
  - POST /api/obligations/extract → AI extraction of obligations

All synchronous OpenAI and Qdrant calls are wrapped in `asyncio.to_thread()`
to prevent blocking the FastAPI event loop under concurrent load.
"""
import io
import re
import asyncio
import functools
import uuid
import traceback
import json
import logging
import html
from datetime import datetime, timezone

from fastapi import APIRouter, UploadFile, File, Form, HTTPException, Depends, Query, Request
from fastapi.responses import StreamingResponse
from pydantic import BaseModel
from supabase import Client
from qdrant_client.http.models import PointStruct, Filter, FieldCondition, MatchValue
from qdrant_client import models
from PyPDF2 import PdfReader
from dotenv import load_dotenv
import os

from app.config import openai_client, qdrant, COLLECTION_NAME, admin_supabase  # CROSS-TENANT: request-path storage uploads still use privileged storage client.
from app.rate_limiter import limiter
from app.task_logger import TaskLogger
from app.token_budget import count_tokens, truncate_to_budget
from app.dependencies import (
    TenantQdrantClient,
    TenantSupabaseClient,
    get_tenant_admin_supabase,  # TenantSupabaseClient factory
    get_tenant_qdrant,
    verify_clerk_token,
    get_tenant_supabase,
)
from app.review_schemas import resolve_contract_status
from app.schemas import ExtractObligationsRequest, ArchiveContractRequest, ConfirmVersionLinkRequest
from app.utils import chunk_text
from app.event_bus import SSEEvent, event_bus
from difflib import SequenceMatcher
from graph import clm_graph
from app.task_logger import TaskLogger
from app.rate_limiter import limiter
from openai import RateLimitError, APITimeoutError, APIConnectionError
import traceback

load_dotenv()

router = APIRouter()
logger = logging.getLogger(__name__)


class ContractPatchRequest(BaseModel):
    title: str | None = None
    end_date: str | None = None
    effective_date: str | None = None
    contract_value: float | None = None
    currency: str | None = None
    jurisdiction: str | None = None
    governing_law: str | None = None


def _sanitize_filename_segment(value: str) -> str:
    cleaned = re.sub(r"[^A-Za-z0-9._-]+", "_", (value or "").strip())
    cleaned = cleaned.strip("._")
    return cleaned or "contract"


def _render_docx_bytes(*, contract_title: str, version_number: int, raw_text: str, finalized_at: str | None) -> bytes:
    from docx import Document

    document = Document()
    document.add_heading(contract_title, level=1)
    meta_line = f"Version {version_number}"
    if finalized_at:
        meta_line = f"{meta_line} — {finalized_at}"
    document.add_paragraph(meta_line)

    paragraphs = [paragraph.strip() for paragraph in raw_text.splitlines() if paragraph.strip()]
    if not paragraphs and raw_text.strip():
        paragraphs = [raw_text.strip()]
    for paragraph in paragraphs:
        document.add_paragraph(paragraph)

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _render_pdf_bytes(*, contract_title: str, version_number: int, raw_text: str, finalized_at: str | None) -> bytes:
    from weasyprint import HTML

    paragraphs = [paragraph.strip() for paragraph in raw_text.splitlines() if paragraph.strip()]
    if not paragraphs and raw_text.strip():
        paragraphs = [raw_text.strip()]
    content_paragraphs = "\n".join(
        f"<p>{html.escape(paragraph)}</p>"
        for paragraph in paragraphs
    )
    meta_line = html.escape(f"Version {version_number}" + (f" — {finalized_at}" if finalized_at else ""))
    html_string = f"""<!doctype html>
<html>
<head>
  <meta charset="utf-8" />
  <style>
    @page {{ size: A4; margin: 2cm; }}
    body {{ font-family: "Times New Roman", serif; color: #111827; }}
    h1 {{ font-size: 16pt; text-align: center; margin-bottom: 8pt; }}
    .meta {{ text-align: center; font-style: italic; color: #4b5563; margin-bottom: 18pt; }}
    p {{ text-align: justify; line-height: 1.5; margin-bottom: 10pt; white-space: pre-wrap; }}
  </style>
</head>
<body>
  <h1>{html.escape(contract_title)}</h1>
  <div class="meta">{meta_line}</div>
  {content_paragraphs}
</body>
</html>"""
    return HTML(string=html_string).write_pdf()


# =====================================================================
# ASYNC WRAPPERS & CALLBACKS
# =====================================================================

def handle_task_result(task: asyncio.Task, contract_id: str, tenant_id: str = None):
    try:
        exc = task.exception()
        if exc:
            # Error is already persisted by TaskLogger inside process_contract_background
            print(f"🚨 [BACKGROUND TASK FAILED] Contract {contract_id}: {exc}")
    except asyncio.CancelledError:
        print(f"⚠️ [BACKGROUND TASK CANCELLED] Contract {contract_id}")

async def async_embed(text: str) -> list[float]:
    response = await asyncio.to_thread(
        openai_client.embeddings.create,
        input=text,
        model="text-embedding-3-small"
    )
    return response.data[0].embedding


async def async_chat_completion(messages: list, response_format=None) -> any:
    kwargs = {"model": "gpt-4o-mini", "messages": messages}
    if response_format:
        kwargs["response_format"] = response_format
    response = await asyncio.to_thread(
        openai_client.chat.completions.create,
        **kwargs
    )
    return response


async def async_qdrant_scroll(collection: str, scroll_filter: Filter, limit: int):
    return await asyncio.to_thread(
        qdrant.scroll,
        collection_name=collection,
        scroll_filter=scroll_filter,
        limit=limit
    )


async def async_clm_graph_invoke(inputs: dict):
    if clm_graph is None:
        raise RuntimeError(
            "FATAL: LangGraph CLM pipeline failed to initialize at startup. "
            "Check graph.py for compilation errors."
        )
    return await asyncio.to_thread(clm_graph.invoke, inputs)


async def async_qdrant_upsert(collection: str, points: list):
    return await asyncio.to_thread(
        qdrant.upsert,
        collection_name=collection,
        points=points
    )


async def publish_contract_event(
    event_type: str,
    tenant_id: str,
    *,
    contract_id: str | None = None,
    data: dict | None = None,
):
    await event_bus.publish(SSEEvent(
        event_type=event_type,
        tenant_id=tenant_id,
        contract_id=contract_id,
        data=data or {},
    ))


def _get_next_version_number(supabase_client: Client, contract_id: str, tenant_id: str) -> int:
    result = supabase_client.table("contract_versions") \
        .select("version_number") \
        .eq("contract_id", contract_id) \
        .eq("tenant_id", tenant_id) \
        .gt("version_number", 0) \
        .order("version_number", desc=True) \
        .limit(1) \
        .execute()
    if result.data:
        return (result.data[0].get("version_number") or 0) + 1
    return 1


def _build_pipeline_snapshot(final_state: dict) -> dict:
    pipeline_snapshot = {}
    for key in [
        "risk_score", "risk_level", "risk_flags_v2", "compliance_findings_v2",
        "draft_revisions_v2", "obligations_v2", "classified_clauses_v2",
        "review_findings", "quick_insights", "banner", "pipeline_output_quality",
        "contract_value", "currency", "end_date", "effective_date",
        "jurisdiction", "governing_law", "counter_proposal",
    ]:
        if key in final_state:
            pipeline_snapshot[key] = final_state[key]
    return pipeline_snapshot


def _extract_pending_file_metadata(version_row: dict) -> dict:
    pipeline_output = version_row.get("pipeline_output") or {}
    return {
        "file_url": pipeline_output.get("_pending_file_path"),
        "file_type": pipeline_output.get("_pending_file_type"),
        "file_size": pipeline_output.get("_pending_file_size"),
    }


def _get_contract_file_metadata(
    supabase_client,
    *,
    contract_id: str,
    tenant_id: str,
) -> dict:
    query = supabase_client.table("contracts") \
        .select("file_url, file_type, file_size") \
        .eq("id", contract_id) \
        .limit(1)
    if not isinstance(supabase_client, TenantSupabaseClient):
        query = query.eq("tenant_id", tenant_id)
    result = query.execute()
    if not result.data:
        return {}
    return result.data[0] or {}


async def _schedule_contract_processing(
    *,
    contract_id: str,
    version_id: str,
    tenant_id: str,
    matter_id: str | None,
    filename: str,
    text_content: str,
):
    try:
        from app.job_queue import enqueue_pipeline

        enqueue_result = await enqueue_pipeline(
            contract_id=contract_id,
            version_id=version_id,
            tenant_id=tenant_id,
            matter_id=matter_id,
            filename=filename,
            text_content=text_content,
        )
        print(f"[QUEUE] Pipeline job enqueued: {enqueue_result['job_id']} for contract {contract_id}")
    except Exception as exc:
        log_id = getattr(exc, "log_id", None)
        print(f"[QUEUE] Failed to enqueue pipeline for {contract_id}: {exc}. Falling back to asyncio.create_task().")
        if log_id is None:
            fallback_logger = TaskLogger(
                tenant_id=tenant_id,
                task_type="pipeline_ingestion",
                contract_id=contract_id,
                version_id=version_id,
                input_metadata={
                    "filename": filename,
                    "version_id": version_id,
                    "matter_id": matter_id,
                    "queue_fallback": True,
                    "queue_error": str(exc),
                },
            )
            log_id = fallback_logger.log_id
        task = asyncio.create_task(
            process_contract_background(
                contract_id=contract_id,
                version_id=version_id,
                tenant_id=tenant_id,
                matter_id=matter_id,
                filename=filename,
                text_content=text_content,
                existing_log_id=log_id,
                task_input_metadata={
                    "queue_fallback": True,
                    "queue_error": str(exc),
                },
            )
        )
        task.add_done_callback(functools.partial(handle_task_result, contract_id=contract_id, tenant_id=tenant_id))


def _create_contract_with_linked_version(
    supabase_client: Client,
    *,
    tenant_id: str,
    matter_id: str | None = None,
    contract_insert_data: dict,
    version_id: str,
    version_insert_data: dict | None = None,
    version_update_data: dict | None = None,
) -> None:
    """
    Avoid circular FK violations by always creating the contract first with
    latest_version_id = NULL, then creating/linking the version row, then
    updating the contract to point at that version.
    """
    if (version_insert_data is None) == (version_update_data is None):
        raise ValueError("Provide exactly one of version_insert_data or version_update_data")
    if not tenant_id:
        raise ValueError("tenant_id is required for tenant-scoped contract writes")

    contract_id = contract_insert_data["id"]
    contract_payload = {
        **contract_insert_data,
        "tenant_id": tenant_id,
        "matter_id": matter_id if matter_id is not None else contract_insert_data.get("matter_id"),
        "latest_version_id": None,
    }
    print(f"DEBUG RLS: Trying to insert contract with tenant_id={tenant_id}")
    print(f"DEBUG MATTER LINK: Trying to insert contract {contract_id} with matter_id={contract_payload.get('matter_id')}")
    supabase_client.table("contracts").insert({**contract_payload, "tenant_id": tenant_id}).execute()

    if version_insert_data is not None:
        version_payload = {
            **version_insert_data,
            "id": version_id,
            "tenant_id": tenant_id,
            "contract_id": contract_id,
        }
        supabase_client.table("contract_versions").insert({**version_payload, "tenant_id": tenant_id}).execute()
    else:
        version_payload = {
            **version_update_data,
            "contract_id": contract_id,
        }
        supabase_client.table("contract_versions").update(version_payload) \
            .eq("id", version_id) \
            .eq("tenant_id", tenant_id) \
            .execute()

    supabase_client.table("contracts").update({
        "latest_version_id": version_id,
    }).eq("id", contract_id).eq("tenant_id", tenant_id).execute()


def _upload_to_storage(
    supabase_client: Client,
    *,
    path: str,
    file_bytes: bytes,
    content_type: str,
) -> str | None:
    try:
        supabase_client.storage.from_("matter-files").upload(
            path=path,
            file=file_bytes,
            file_options={"content-type": content_type},
        )
        return path
    except Exception as storage_err:
        logger.error("🚨 [Storage] Upload FAILED for path=%s: %s: %s", path, type(storage_err).__name__, storage_err)
        return None


def _build_pipeline_input_metadata(
    *,
    version_id: str,
    matter_id: str | None,
    filename: str,
    text_content: str,
    extra_metadata: dict | None = None,
) -> tuple[str, bool, dict | None, dict]:
    raw_tokens = count_tokens(text_content, "gpt-4o-mini")
    safe_document, is_truncated, original_token_count = truncate_to_budget(
        text_content,
        max_tokens=80_000,
        model="gpt-4o-mini",
        strategy="tail_preserve",
    )

    truncation_warning = None
    if is_truncated:
        truncation_warning = {
            "original_tokens": original_token_count,
            "truncated_to": 80_000,
            "chars_removed": max(0, len(text_content) - len(safe_document)),
            "strategy": "tail_preserve",
            "message": (
                f"Document was {original_token_count:,} tokens and was truncated to 80,000 tokens "
                "using a tail-preserve strategy. The beginning and end of the document were retained."
            ),
        }
        logger.warning(
            "[PIPELINE] Contract %s truncated from %s to 80,000 tokens before analysis.",
            version_id,
            f"{original_token_count:,}",
        )

    input_metadata = {
        "filename": filename,
        "raw_tokens": raw_tokens,
        "is_truncated": is_truncated,
        "version_id": version_id,
        "matter_id": matter_id,
        **(extra_metadata or {}),
    }
    return safe_document, is_truncated, truncation_warning, input_metadata


# =====================================================================
# ENDPOINTS
# =====================================================================

@router.get("/contracts")
@limiter.limit("60/minute")
async def list_contracts(
    request: Request,
    tab: str = "Archived",
    limit: int | None = None,
    sort_by: str = "created_at",
    sort_order: str = "desc",
    claims: dict = Depends(verify_clerk_token),
    supabase: Client = Depends(get_tenant_supabase),
):
    """
    GET /api/v1/contracts?tab=Archived|active|Active Contracts|templates
    Returns contracts filtered by status category for the authenticated tenant.
    """
    tenant_id = claims["verified_tenant_id"]
    tab_lower = tab.lower().strip()
    sort_by_lower = sort_by.lower().strip()
    sort_order_lower = sort_order.lower().strip()

    if limit is not None and limit < 1:
        raise HTTPException(status_code=400, detail="limit must be greater than 0")
    if sort_by_lower not in {"created_at", "updated_at"}:
        raise HTTPException(status_code=400, detail="sort_by must be created_at or updated_at")
    if sort_order_lower not in {"asc", "desc"}:
        raise HTTPException(status_code=400, detail="sort_order must be asc or desc")

    try:
        query = supabase.table("contracts").select("*").eq("tenant_id", tenant_id)

        if tab_lower in ("archived", "expired", "terminated"):
            query = query.in_("status", ["EXPIRED", "TERMINATED", "ARCHIVED", "Superseded"])
        elif tab_lower in ("active", "active contracts"):
            # Use an exclusion-based filter so newly introduced in-flight statuses
            # like "Processing" still appear in the primary documents view.
            query = query \
                .neq("status", "ARCHIVED") \
                .neq("status", "EXPIRED") \
                .neq("status", "TERMINATED") \
                .neq("status", "Superseded") \
                .neq("status", "TEMPLATE")
        elif tab_lower in ("templates", "templates & playbooks"):
            query = query.eq("status", "TEMPLATE")
        else:
            # Default: exclude ARCHIVED (and other terminal statuses) from global queries
            query = query \
                .neq("status", "ARCHIVED") \
                .neq("status", "EXPIRED") \
                .neq("status", "TERMINATED") \
                .neq("status", "Superseded")

        query = query.order(sort_by_lower, desc=sort_order_lower == "desc")
        if limit is not None:
            query = query.limit(limit)
        result = query.execute()

        return {"data": result.data or []}
    except Exception as e:
        print(f"[GET /contracts] Error: {e}")
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=str(e))


@router.get("/contracts/{contract_id}")
@limiter.limit("60/minute")
async def get_contract(
    request: Request,
    contract_id: str,
    claims: dict = Depends(verify_clerk_token),
    supabase: Client = Depends(get_tenant_supabase)
):
    """
    GET /api/contracts/{contract_id}
    Returns a single contract for the authenticated tenant.
    Related AI extraction tables are intentionally not joined here so a
    partially processed contract can still load its detail page cleanly.
    """
    tenant_id = claims["verified_tenant_id"]

    try:
        result = supabase.table("contracts") \
            .select("*") \
            .eq("id", contract_id) \
            .eq("tenant_id", tenant_id) \
            .limit(1) \
            .execute()

        if not result.data:
            raise HTTPException(status_code=404, detail="Contract not found")

        return {"data": result.data[0]}
    except HTTPException:
        raise
    except Exception as e:
        print(f"[GET /contracts/{contract_id}] Error: {e}")
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=str(e))


@router.get("/contracts/{contract_id}/versions/{version_id}/export")
@limiter.limit("30/minute")
async def export_contract_version(
    request: Request,
    contract_id: str,
    version_id: str,
    format: str = Query(..., pattern="^(docx|pdf)$"),
    claims: dict = Depends(verify_clerk_token),
    supabase: Client = Depends(get_tenant_supabase),
):
    try:
        tenant_id = claims["verified_tenant_id"]
        if format not in {"docx", "pdf"}:
            raise HTTPException(status_code=400, detail="Invalid export format")

        contract_res = supabase.table("contracts").select("id, title") \
            .eq("id", contract_id) \
            .eq("tenant_id", tenant_id) \
            .limit(1) \
            .execute()
        if not contract_res.data:
            raise HTTPException(status_code=404, detail="Contract not found")
        contract = contract_res.data[0]

        try:
            version_res = supabase.table("contract_versions").select(
                "id, contract_id, version_number, raw_text, uploaded_filename, finalized_at, created_at"
            ) \
                .eq("id", version_id) \
                .eq("contract_id", contract_id) \
                .eq("tenant_id", tenant_id) \
                .limit(1) \
                .execute()
        except Exception as exc:
            if "contract_versions.finalized_at" not in str(exc):
                raise
            version_res = supabase.table("contract_versions").select(
                "id, contract_id, version_number, raw_text, uploaded_filename, created_at"
            ) \
                .eq("id", version_id) \
                .eq("contract_id", contract_id) \
                .eq("tenant_id", tenant_id) \
                .limit(1) \
                .execute()
        if not version_res.data:
            raise HTTPException(status_code=404, detail="Contract version not found")
        version = version_res.data[0]

        contract_title = str(contract.get("title") or version.get("uploaded_filename") or "Contract")
        version_number = int(version.get("version_number") or 0)
        raw_text = str(version.get("raw_text") or "")
        finalized_at = version.get("finalized_at") or version.get("created_at")
        finalized_label = None
        if finalized_at:
            try:
                finalized_label = datetime.fromisoformat(str(finalized_at).replace("Z", "+00:00")) \
                    .astimezone(timezone.utc) \
                    .strftime("%Y-%m-%d %H:%M UTC")
            except ValueError:
                finalized_label = str(finalized_at)

        if format == "docx":
            file_bytes = _render_docx_bytes(
                contract_title=contract_title,
                version_number=version_number,
                raw_text=raw_text,
                finalized_at=finalized_label,
            )
            media_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            extension = "docx"
        else:
            file_bytes = _render_pdf_bytes(
                contract_title=contract_title,
                version_number=version_number,
                raw_text=raw_text,
                finalized_at=finalized_label,
            )
            media_type = "application/pdf"
            extension = "pdf"

        filename = (
            f"{_sanitize_filename_segment(contract_title)}"
            f"_V{version_number}_{datetime.now(timezone.utc).strftime('%Y%m%dT%H%M%SZ')}.{extension}"
        )

        return StreamingResponse(
            io.BytesIO(file_bytes),
            media_type=media_type,
            headers={"Content-Disposition": f'attachment; filename="{filename}"'},
        )
    except HTTPException:
        raise
    except Exception as e:
        print(f"[EXPORT /contracts/{contract_id}/versions/{version_id}] Error: {e}")
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=str(e))


@router.patch("/contracts/{contract_id}")
@limiter.limit("60/minute")
async def update_contract(
    request: Request,
    contract_id: str,
    req: ContractPatchRequest,
    claims: dict = Depends(verify_clerk_token),
    supabase: Client = Depends(get_tenant_supabase),
):
    tenant_id = claims["verified_tenant_id"]
    payload = req.dict(exclude_unset=True)
    if not payload:
        raise HTTPException(status_code=400, detail="No valid fields to update")

    try:
        result = (
            supabase.table("contracts")
            .update(payload)
            .eq("id", contract_id)
            .eq("tenant_id", tenant_id)
            .select()
            .single()
            .execute()
        )
        if not result.data:
            raise HTTPException(status_code=404, detail="Contract not found")

        return {"contract": result.data}
    except HTTPException:
        raise
    except Exception as e:
        print(f"[PATCH /contracts/{contract_id}] Error: {e}")
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=str(e))

async def _execute_pipeline_and_persist(
    *,
    contract_id: str,
    version_id: str,
    tenant_id: str,
    tenant_sb: TenantSupabaseClient,
    matter_id: str | None,
    filename: str,
    text_content: str,
    safe_document: str,
    is_truncated: bool,
    truncation_warning: dict | None,
    task_logger: TaskLogger,
) -> dict:
    print(f"🔄 [LANGGRAPH] Invoking agentic workflow for document ID: {contract_id}")

    final_state = await async_clm_graph_invoke({
        "contract_id": contract_id,
        "raw_document": safe_document,
        "_task_logger": task_logger,
        "_event_bus": event_bus,
        "_tenant_id": tenant_id,
    })

    print(f"[LANGGRAPH] Execution complete for {contract_id}. Raw JSON output:")
    print(json.dumps({
        "risk_score": final_state.get("risk_score"),
        "contract_value": final_state.get("contract_value"),
        "end_date": final_state.get("end_date"),
        "currency": final_state.get("currency"),
        "effective_date": final_state.get("effective_date")
    }, default=str, indent=2))

    score = final_state.get("risk_score", 0.0)
    risk_level = final_state.get("risk_level") or (
        "High" if score >= 75.0 else ("Medium" if score >= 40.0 else ("Low" if score > 0 else "Safe"))
    )

    draft_revisions_payload = {
        "latest_text": text_content,
        "findings": final_state.get("draft_revisions", []),
        "history": [],
    }
    if truncation_warning:
        draft_revisions_payload["truncation_warning"] = truncation_warning

    pipeline_quality = final_state.get("pipeline_output_quality", "complete")
    sentinel_findings = [
        finding for finding in (final_state.get("review_findings") or [])
        if finding.get("is_sentinel")
    ]
    contract_status = resolve_contract_status(pipeline_quality)
    if pipeline_quality == "empty":
        logger.warning(
            "pipeline_output_empty | contract_id=%s | tenant_id=%s | sentinel_count=%s",
            contract_id,
            tenant_id,
            len(sentinel_findings),
        )
    elif pipeline_quality == "partial":
        logger.info(
            "pipeline_output_partial | contract_id=%s | tenant_id=%s | sentinel_count=%s",
            contract_id,
            tenant_id,
            len(sentinel_findings),
        )

    update_payload = {
        "status": contract_status,
        "contract_value": float(final_state.get("contract_value", 0.0) or 0.0),
        "end_date": final_state.get("end_date", "Unknown"),
        "effective_date": final_state.get("effective_date", None),
        "jurisdiction": final_state.get("jurisdiction", None),
        "governing_law": final_state.get("governing_law", None),
        "risk_level": risk_level,
        "currency": final_state.get("currency", "IDR"),
        "counter_proposal": final_state.get("counter_proposal"),
        "is_truncated": is_truncated,
        "draft_revisions": draft_revisions_payload,
    }
    update_payload.pop("file_url", None)
    update_payload.pop("file_path", None)
    dropped_none_keys = [key for key, value in update_payload.items() if value is None]
    update_payload = {key: value for key, value in update_payload.items() if value is not None}

    print(f"[SUPABASE UPDATE] Updating contract_id: {contract_id} with pipeline results.")
    print(f"[DEBUG] Update payload keys: {list(update_payload.keys())}")
    print(f"[DEBUG] file_url in payload: {'file_url' in update_payload}")
    print(f"[DEBUG] file_path in payload: {'file_path' in update_payload}")
    if dropped_none_keys:
        print(f"[DEBUG] Dropped None-valued keys from update payload: {dropped_none_keys}")
    print(json.dumps(update_payload, indent=2, default=str))

    try:
        existing_file_meta = _get_contract_file_metadata(
            tenant_sb,
            contract_id=contract_id,
            tenant_id=tenant_id,
        )
        res = tenant_sb.table("contracts").update(update_payload).eq("id", contract_id).execute()
        updated_contract = (res.data or [{}])[0]
        existing_file_url = existing_file_meta.get("file_url")
        if existing_file_url and not updated_contract.get("file_url"):
            print(
                f"⚠️ [SUPABASE UPDATE] file_url unexpectedly missing after pipeline update for {contract_id}. "
                "Restoring previously saved storage path."
            )
            restore_payload = {"file_url": existing_file_url}
            if existing_file_meta.get("file_type") is not None:
                restore_payload["file_type"] = existing_file_meta.get("file_type")
            if existing_file_meta.get("file_size") is not None:
                restore_payload["file_size"] = existing_file_meta.get("file_size")
            restore_res = tenant_sb.table("contracts").update(restore_payload).eq("id", contract_id).execute()
            print(f"[SUPABASE UPDATE] Restored file metadata: {restore_res.data}")
        print(f"[SUPABASE UPDATE] Success! Response: {res.data}")
    except Exception as exc:
        print(f"!!! [SUPABASE UPDATE ERROR] {exc}")
        traceback.print_exc()

    try:
        pipeline_snapshot = _build_pipeline_snapshot(final_state)
        if truncation_warning:
            pipeline_snapshot["truncation_warning"] = truncation_warning
        tenant_sb.table("contract_versions").update({
            "raw_text": text_content[:500000],
            "pipeline_output": pipeline_snapshot,
            "risk_score": float(score),
            "risk_level": risk_level,
            "uploaded_filename": filename,
            "is_truncated": is_truncated,
        }).eq("id", version_id).execute()
        print(f"✅ [War Room] Updated version snapshot id={version_id} for contract {contract_id}")
    except Exception as ver_err:
        print(f"⚠️ [War Room] Failed to persist version snapshot: {ver_err}")
        traceback.print_exc()

    obligations = final_state.get("extracted_obligations", [])
    if obligations:
        obligations_data = [
            {
                "tenant_id": tenant_id,
                "contract_id": contract_id,
                "description": ob.get("description", ""),
                "due_date": ob.get("due_date"),
                "status": "pending",
            }
            for ob in obligations if ob.get("description")
        ]
        if obligations_data:
            try:
                tenant_sb.table("contract_obligations").insert(obligations_data).execute()
                print(f"✅ Inserted {len(obligations_data)} obligations for contract {contract_id}")
            except Exception as ob_err:
                print(f"⚠️ Failed to insert obligations: {ob_err}")

    clauses = final_state.get("classified_clauses", [])
    if clauses:
        clauses_data = [
            {
                "tenant_id": tenant_id,
                "contract_id": contract_id,
                "clause_type": cl.get("clause_type", "Other"),
                "original_text": cl.get("original_text", ""),
                "ai_summary": cl.get("ai_summary"),
            }
            for cl in clauses if cl.get("original_text")
        ]
        if clauses_data:
            try:
                tenant_sb.table("contract_clauses").insert(clauses_data).execute()
                print(f"✅ Inserted {len(clauses_data)} clauses for contract {contract_id}")
            except Exception as cl_err:
                print(f"⚠️ Failed to insert clauses: {cl_err}")

    chunks = chunk_text(text_content, chunk_size=1500, overlap=200)
    print(f"🔄 [BACKGROUND] Document split into {len(chunks)} chunks. Starting Qdrant vectorization...")

    points = []
    for i, chunk in enumerate(chunks):
        chunk_embed = await async_embed(chunk)
        page_match = re.search(r'\[Page (\d+)\]', chunk)
        page_number = page_match.group(1) if page_match else "Unknown"

        points.append(
            PointStruct(
                id=str(uuid.uuid4()),
                vector=chunk_embed,
                payload={
                    "tenant_id": tenant_id,
                    "contract_id": contract_id,
                    "chunk_index": i,
                    "text": chunk,
                    "page_number": page_number,
                },
            )
        )

    if points:
        batch_size = 100
        for i in range(0, len(points), batch_size):
            batch = points[i:i + batch_size]
            await async_qdrant_upsert(COLLECTION_NAME, batch)

    findings_count = len(final_state.get("compliance_findings_v2", [])) + len(final_state.get("risk_flags_v2", []))
    return {
        "contract_status": contract_status,
        "pipeline_output_quality": pipeline_quality,
        "sentinel_count": len(sentinel_findings),
        "risk_score": final_state.get("risk_score", 0),
        "risk_level": risk_level,
        "findings_count": findings_count,
        "obligations_count": len(final_state.get("obligations_v2", [])),
        "clauses_count": len(final_state.get("classified_clauses_v2", [])),
        "chunks_vectorized": len(chunks),
        "version_id": version_id,
        "filename": filename,
    }


async def process_contract_background(
    contract_id: str,
    version_id: str,
    tenant_id: str,
    matter_id: str | None,
    filename: str,
    text_content: str,
    attempt_number: int = 1,
    parent_task_id: str = None,
    existing_log_id: str | None = None,
    task_input_metadata: dict | None = None,
):
    print(f"🚀 [BACKGROUND] Starting LangGraph pipeline for contract_id={contract_id}, version_id={version_id}")
    tenant_sb = get_tenant_admin_supabase(tenant_id)  # TenantSupabaseClient

    contract_check = tenant_sb.table("contracts") \
        .select("id, status") \
        .eq("id", contract_id) \
        .limit(1) \
        .execute()
    if not contract_check.data:
        print(f"[PIPELINE ABORTED] Contract {contract_id} no longer exists.")
        return {"status": "aborted", "reason": "missing_contract"}

    current_status = contract_check.data[0].get("status")
    if current_status == "ARCHIVED":
        print(f"[PIPELINE ABORTED] Contract {contract_id} has been archived.")
        return {"status": "aborted", "reason": "archived"}

    safe_document, is_truncated, truncation_warning, input_metadata = _build_pipeline_input_metadata(
        version_id=version_id,
        matter_id=matter_id,
        filename=filename,
        text_content=text_content,
        extra_metadata=task_input_metadata,
    )

    task_logger = TaskLogger(
        tenant_id=tenant_id,
        task_type="pipeline_ingestion",
        contract_id=contract_id,
        version_id=version_id,
        input_metadata=input_metadata,
        attempt_number=attempt_number,
        parent_task_id=parent_task_id,
        existing_log_id=existing_log_id,
    )

    try:
        if current_status != "Processing":
            tenant_sb.table("contracts").update({
                "status": "Processing",
            }).eq("id", contract_id).execute()
            await publish_contract_event(
                "contract.status_changed",
                tenant_id,
                contract_id=contract_id,
                data={
                    "contract_id": contract_id,
                    "contract_title": filename,
                    "old_status": current_status,
                    "new_status": "Processing",
                    "message": f"{filename} is processing",
                },
            )

        await publish_contract_event(
            "pipeline.started",
            tenant_id,
            contract_id=contract_id,
            data={
                "filename": filename,
                "text_length": len(text_content),
                "total_agents": 8,
                "message": "AI contract analysis started",
            },
        )

        result_summary = await _execute_pipeline_and_persist(
            contract_id=contract_id,
            version_id=version_id,
            tenant_id=tenant_id,
            tenant_sb=tenant_sb,
            matter_id=matter_id,
            filename=filename,
            text_content=text_content,
            safe_document=safe_document,
            is_truncated=is_truncated,
            truncation_warning=truncation_warning,
            task_logger=task_logger,
        )

        completion_event_type = (
            "pipeline.completed_incomplete"
            if result_summary["contract_status"] == "Review_Incomplete"
            else "pipeline.completed"
        )
        completion_message = (
            "AI contract analysis completed with incomplete output"
            if result_summary["contract_status"] == "Review_Incomplete"
            else "AI contract analysis completed"
        )
        await publish_contract_event(
            completion_event_type,
            tenant_id,
            contract_id=contract_id,
            data={
                "risk_score": result_summary["risk_score"],
                "risk_level": result_summary["risk_level"],
                "findings_count": result_summary["findings_count"],
                "quality": result_summary["pipeline_output_quality"],
                "sentinel_count": result_summary["sentinel_count"],
                "message": completion_message,
            },
        )
        await publish_contract_event(
            "contract.risk_updated",
            tenant_id,
            contract_id=contract_id,
            data={
                "risk_score": result_summary["risk_score"],
                "risk_level": result_summary["risk_level"],
            },
        )
        await publish_contract_event(
            "contract.status_changed",
            tenant_id,
            contract_id=contract_id,
            data={
                "contract_id": contract_id,
                "contract_title": filename,
                "old_status": "Processing",
                "new_status": result_summary["contract_status"],
                "message": (
                    f"{filename} review is incomplete"
                    if result_summary["contract_status"] == "Review_Incomplete"
                    else f"{filename} has been reviewed"
                ),
            },
        )

        task_logger.complete(result_summary=result_summary)
        print(f"[BACKGROUND] process_contract_background successfully completed for {contract_id}.")
        return result_summary
    except Exception as exc:
        task_logger.fail(exc)
        tenant_sb.table("contracts").update({
            "status": "Failed",
            "draft_revisions": {
                "error_log_id": task_logger.log_id,
                "error_summary": str(exc)[:500],
            }
        }).eq("id", contract_id).execute()
        await publish_contract_event(
            "pipeline.failed",
            tenant_id,
            contract_id=contract_id,
            data={
                "error": str(exc)[:500],
                "message": "AI contract analysis failed",
            },
        )
        await publish_contract_event(
            "contract.status_changed",
            tenant_id,
            contract_id=contract_id,
            data={
                "contract_id": contract_id,
                "contract_title": filename,
                "old_status": "Processing",
                "new_status": "Failed",
                "message": f"{filename} failed to process",
            },
        )
        print(f"!!! [BACKGROUND] Unhandled Exception during process_contract_background: {exc}")
        traceback.print_exc()
        raise

TRANSIENT_ERRORS = (RateLimitError, APITimeoutError, APIConnectionError, ConnectionError, TimeoutError)
MAX_RETRY_ATTEMPTS = 3
BASE_DELAY_SECONDS = 5  # 5s, 10s, 20s exponential backoff

async def process_contract_with_retry(contract_id, version_id, tenant_id, matter_id, filename, text_content):
    """Wrapper that retries process_contract_background on transient failures."""
    
    parent_task_id = None
    tenant_sb = get_tenant_admin_supabase(tenant_id)  # TenantSupabaseClient
    
    for attempt in range(1, MAX_RETRY_ATTEMPTS + 1):
        try:
            await process_contract_background(
                contract_id=contract_id,
                version_id=version_id,
                tenant_id=tenant_id,
                matter_id=matter_id,
                filename=filename,
                text_content=text_content,
                attempt_number=attempt,
                parent_task_id=parent_task_id,
            )
            return  # Success — exit retry loop
            
        except TRANSIENT_ERRORS as e:
            
            if attempt < MAX_RETRY_ATTEMPTS:
                # Update contract status to show retry
                tenant_sb.table("contracts").update({
                    "status": f"Retrying ({attempt}/{MAX_RETRY_ATTEMPTS})"
                }).eq("id", contract_id).execute()
                await publish_contract_event(
                    "contract.status_changed",
                    tenant_id,
                    contract_id=contract_id,
                    data={
                        "contract_id": contract_id,
                        "old_status": "Processing",
                        "new_status": f"Retrying ({attempt}/{MAX_RETRY_ATTEMPTS})",
                        "message": f"Retry attempt {attempt} started",
                    },
                )
                
                delay = BASE_DELAY_SECONDS * (2 ** (attempt - 1))  # 5, 10, 20 seconds
                await asyncio.sleep(delay)
            else:
                # All retries exhausted
                raise
        
        except Exception as e:
            # Non-transient error — do NOT retry
            raise

@router.post("/upload")
@limiter.limit("5/minute")
@limiter.limit("30/hour")
async def upload_contract(
    request: Request,
    file: UploadFile = File(...),
    matter_id: str = Form(None),
    contract_id: str = Form(None),
    parent_contract_id: str = Form(None),
    document_category: str = Form("Uncategorized"),
    parent_id: str = Form(None),
    relationship_type: str = Form(None),
    claims: dict = Depends(verify_clerk_token),
    supabase: Client = Depends(get_tenant_supabase)
):
    tenant_id = claims["verified_tenant_id"]
    matter_id = matter_id.strip() if isinstance(matter_id, str) else matter_id
    if matter_id == "":
        matter_id = None

    if matter_id:
        matter_check = supabase.table("matters").select("tenant_id").eq("id", matter_id).eq("tenant_id", tenant_id).execute()
        if matter_check.data and matter_check.data[0].get("tenant_id"):
            tenant_id = matter_check.data[0]["tenant_id"]

    print(f"[ENDPOINT HIT] /api/upload called for filename: {file.filename} (Resolved Tenant: {tenant_id})")

    if not file.filename.lower().endswith('.pdf'):
        raise HTTPException(status_code=400, detail="Hanya menerima file berformat PDF.")

    contents = await file.read()

    if not contents.startswith(b'%PDF-'):
        raise HTTPException(status_code=403, detail="Peringatan Keamanan: File ini mencoba menyamar sebagai PDF. Ditolak.")

    try:
        import fitz
        import pymupdf4llm
        doc = fitz.open(stream=contents, filetype="pdf")
        text_content = pymupdf4llm.to_markdown(doc)

        if not text_content.strip():
            raise HTTPException(status_code=400, detail="Kami tidak dapat membaca dokumen ini. Pastikan PDF tidak dienkripsi atau berupa gambar hasil scan tanpa OCR.")

        storage_key_id = parent_contract_id or contract_id or str(uuid.uuid4())
        safe_name = re.sub(r'[^a-zA-Z0-9.-]', '_', file.filename or 'document.pdf')
        file_path = f"{tenant_id}/{matter_id}/{storage_key_id[:8]}_{safe_name}"
        file_content_type = file.content_type or "application/pdf"
        file_size = len(contents)

        file_path = _upload_to_storage(
            admin_supabase,  # CROSS-TENANT: storage API is outside the table wrapper; tenant isolation is enforced by the storage path.
            path=file_path,
            file_bytes=contents,
            content_type=file_content_type,
        )
        if file_path:
            print(f"✅ [Storage] Uploaded {file_path} ({file_size} bytes)")

        if parent_contract_id:
            existing_contract = supabase.table("contracts") \
                .select("id, status, title") \
                .eq("id", parent_contract_id) \
                .eq("tenant_id", tenant_id) \
                .limit(1) \
                .execute()
            if not existing_contract.data:
                raise HTTPException(status_code=404, detail="Parent contract not found.")

            previous_status = existing_contract.data[0].get("status")
            next_version = _get_next_version_number(supabase, parent_contract_id, tenant_id)
            version_id = str(uuid.uuid4())

            supabase.table("contract_versions").insert({
                "id": version_id,
                "tenant_id": tenant_id,
                "contract_id": parent_contract_id,
                "version_number": next_version,
                "raw_text": text_content[:500000],
                "uploaded_filename": file.filename,
                "pipeline_output": {},
            }).execute()

            update_data = {
                "status": "Queued",
                "version_count": next_version,
                "latest_version_id": version_id,
            }
            if file_path:
                update_data["file_url"] = file_path
                update_data["file_type"] = file_content_type
                update_data["file_size"] = file_size
            supabase.table("contracts") \
                .update(update_data) \
                .eq("id", parent_contract_id) \
                .eq("tenant_id", tenant_id) \
                .execute()
            print(f"✅ [DB] Updated parent contract {parent_contract_id} and created version {version_id}")
            await publish_contract_event(
                "contract.status_changed",
                tenant_id,
                contract_id=parent_contract_id,
                data={
                    "contract_id": parent_contract_id,
                    "contract_title": file.filename,
                    "old_status": previous_status,
                    "new_status": "Queued",
                    "message": f"{file.filename} is queued for AI processing",
                },
            )
            await _schedule_contract_processing(
                contract_id=parent_contract_id,
                version_id=version_id,
                tenant_id=tenant_id,
                matter_id=matter_id,
                filename=file.filename,
                text_content=text_content,
            )
            return {
                "status": "success",
                "contract_id": parent_contract_id,
                "is_version_candidate": False,
                "message": "Versi baru diunggah dan diproses di latar belakang.",
            }

        best_match = None
        best_score = 0.0
        try:
            contracts_query = supabase.table("contracts") \
                .select("id, title, matter_id") \
                .eq("tenant_id", tenant_id) \
                .neq("status", "ARCHIVED")
            if matter_id:
                contracts_query = contracts_query.eq("matter_id", matter_id)
            existing_contracts = contracts_query.execute()

            upload_name = (file.filename or "").lower().replace(".pdf", "").strip()
            for existing_contract in existing_contracts.data or []:
                existing_title = (existing_contract.get("title") or "").lower().replace(".pdf", "").strip()
                if not existing_title:
                    continue
                score = SequenceMatcher(None, upload_name, existing_title).ratio()
                if score >= 0.6 and score > best_score:
                    best_score = score
                    best_match = existing_contract
        except Exception as match_err:
            print(f"⚠️ [War Room] Fuzzy match check failed (non-fatal): {match_err}")

        if best_match:
            pending_version_id = str(uuid.uuid4())
            supabase.table("contract_versions").insert({
                "id": pending_version_id,
                "tenant_id": tenant_id,
                "contract_id": best_match["id"],
                "version_number": 0,
                "raw_text": text_content[:500000],
                "uploaded_filename": file.filename,
                "pipeline_output": {
                    "_pending": True,
                    "_pending_file_path": file_path,
                    "_pending_file_type": file_content_type,
                    "_pending_file_size": file_size,
                    "_pending_matter_id": matter_id,
                },
            }).execute()

            return {
                "status": "pending_version_decision",
                "pending_version_id": pending_version_id,
                "is_version_candidate": True,
                "matched_contract_id": best_match["id"],
                "matched_contract_title": best_match.get("title", ""),
                "similarity_score": round(best_score, 3),
                "filename": file.filename,
                "message": "Dokumen ini mungkin merupakan versi baru dari kontrak yang sudah ada. Silakan konfirmasi.",
            }

        contract_id = contract_id or str(uuid.uuid4())
        version_id = str(uuid.uuid4())

        insert_data = {
            "id": contract_id,
            "tenant_id": tenant_id,
            "matter_id": matter_id,
            "title": file.filename,
            "file_type": file_content_type,
            "file_size": file_size,
            "document_category": document_category or "Uncategorized",
            "status": "Queued",
            "version_count": 1,
        }
        if file_path:
            insert_data["file_url"] = file_path

        _create_contract_with_linked_version(
            supabase,
            tenant_id=tenant_id,
            matter_id=matter_id,
            contract_insert_data=insert_data,
            version_id=version_id,
            version_insert_data={
            "tenant_id": tenant_id,
            "version_number": 1,
            "raw_text": text_content[:500000],
            "uploaded_filename": file.filename,
            "pipeline_output": {},
            },
        )

        if parent_id:
            try:
                supabase.table("document_relationships").insert({
                    "tenant_id": tenant_id,
                    "parent_id": parent_id,
                    "child_id": contract_id,
                    "relationship_type": relationship_type or "related_to",
                }).execute()
            except Exception as rel_err:
                print(f"⚠️ [Genealogy] Insert failed (non-fatal): {rel_err}")

        await publish_contract_event(
            "contract.created",
            tenant_id,
            contract_id=contract_id,
            data={
                "contract_id": contract_id,
                "contract_title": file.filename,
                "status": "Queued",
                "matter_id": matter_id,
                "message": f"{file.filename} uploaded",
            },
        )
        await publish_contract_event(
            "contract.status_changed",
            tenant_id,
            contract_id=contract_id,
            data={
                "contract_id": contract_id,
                "contract_title": file.filename,
                "old_status": None,
                "new_status": "Queued",
                "message": f"{file.filename} is queued for AI processing",
            },
        )

        await _schedule_contract_processing(
            contract_id=contract_id,
            version_id=version_id,
            tenant_id=tenant_id,
            matter_id=matter_id,
            filename=file.filename,
            text_content=text_content,
        )
        print(f"[ENDPOINT SUCCESS] Background task scheduled for contract_id: {contract_id}. Returning 200 immediately.")

        return {
            "status": "success",
            "message": "Upload diproses di latar belakang.",
            "contract_id": contract_id,
            "is_version_candidate": False,
            "smart_metadata": {},
        }
    except Exception as e:
        print(f"API Upload Error: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@router.post("/upload/confirm-version")
@limiter.limit("60/minute")
async def confirm_version_link(
    request: Request,
    claims: dict = Depends(verify_clerk_token),
    supabase: Client = Depends(get_tenant_supabase),
):
    try:
        base_tenant_id = claims["verified_tenant_id"]

        content_type = request.headers.get("content-type", "")
        if "application/json" in content_type:
            payload = await request.json()
        else:
            payload = dict(await request.form())

        pending_version_id = payload.get("pending_version_id")
        matched_contract_id = payload.get("matched_contract_id")
        action = payload.get("action")
        matter_id = payload.get("matter_id")

        if pending_version_id and matched_contract_id and action:
            parent_res = supabase.table("contracts") \
                .select("id, title, version_count, matter_id, tenant_id, status") \
                .eq("id", matched_contract_id) \
                .limit(1) \
                .execute()
            if not parent_res.data:
                raise HTTPException(status_code=404, detail="Matched contract not found.")

            parent = parent_res.data[0]
            tenant_id = parent.get("tenant_id") or base_tenant_id
            pending_res = supabase.table("contract_versions") \
                .select("*") \
                .eq("id", pending_version_id) \
                .eq("tenant_id", tenant_id) \
                .limit(1) \
                .execute()
            if not pending_res.data:
                raise HTTPException(status_code=404, detail="Pending version not found.")

            pending_version = pending_res.data[0]
            if not (pending_version.get("pipeline_output") or {}).get("_pending"):
                raise HTTPException(status_code=400, detail="Version candidate has already been resolved.")
            text_content = pending_version.get("raw_text", "")
            filename = pending_version.get("uploaded_filename") or "uploaded.pdf"
            file_meta = _extract_pending_file_metadata(pending_version)
            resolved_matter_id = matter_id or pending_version.get("pipeline_output", {}).get("_pending_matter_id") or parent.get("matter_id")

            if action == "confirm":
                next_version = _get_next_version_number(supabase, matched_contract_id, tenant_id)
                supabase.table("contract_versions").update({
                    "contract_id": matched_contract_id,
                    "version_number": next_version,
                    "pipeline_output": {},
                }).eq("id", pending_version_id).eq("tenant_id", tenant_id).execute()

                contract_update = {
                    "status": "Queued",
                    "version_count": next_version,
                    "latest_version_id": pending_version_id,
                }
                if file_meta.get("file_url"):
                    contract_update["file_url"] = file_meta["file_url"]
                    contract_update["file_type"] = file_meta.get("file_type")
                    contract_update["file_size"] = file_meta.get("file_size")
                supabase.table("contracts").update(contract_update) \
                    .eq("id", matched_contract_id) \
                    .eq("tenant_id", tenant_id) \
                    .execute()

                await publish_contract_event(
                    "contract.status_changed",
                    tenant_id,
                    contract_id=matched_contract_id,
                    data={
                        "contract_id": matched_contract_id,
                        "contract_title": parent.get("title") or filename,
                        "old_status": parent.get("status"),
                        "new_status": "Queued",
                        "message": f"{filename} is queued as a new version",
                    },
                )
                target_contract_id = matched_contract_id
            elif action == "reject":
                target_contract_id = str(uuid.uuid4())
                contract_insert = {
                    "id": target_contract_id,
                    "tenant_id": tenant_id,
                    "matter_id": resolved_matter_id,
                    "title": filename,
                    "status": "Queued",
                    "version_count": 1,
                }
                if file_meta.get("file_url"):
                    contract_insert["file_url"] = file_meta["file_url"]
                    contract_insert["file_type"] = file_meta.get("file_type")
                    contract_insert["file_size"] = file_meta.get("file_size")

                _create_contract_with_linked_version(
                    supabase,
                    tenant_id=tenant_id,
                    matter_id=resolved_matter_id,
                    contract_insert_data=contract_insert,
                    version_id=pending_version_id,
                    version_update_data={
                        "version_number": 1,
                        "pipeline_output": {},
                    },
                )

                await publish_contract_event(
                    "contract.created",
                    tenant_id,
                    contract_id=target_contract_id,
                    data={
                        "contract_id": target_contract_id,
                        "contract_title": filename,
                        "status": "Queued",
                        "matter_id": resolved_matter_id,
                        "message": f"{filename} uploaded",
                    },
                )
                await publish_contract_event(
                    "contract.status_changed",
                    tenant_id,
                    contract_id=target_contract_id,
                    data={
                        "contract_id": target_contract_id,
                        "contract_title": filename,
                        "old_status": None,
                        "new_status": "Queued",
                        "message": f"{filename} is queued for AI processing",
                    },
                )
            else:
                raise HTTPException(status_code=400, detail=f"Invalid action: {action}. Must be 'confirm' or 'reject'.")

            await _schedule_contract_processing(
                contract_id=target_contract_id,
                version_id=pending_version_id,
                tenant_id=tenant_id,
                matter_id=resolved_matter_id,
                filename=filename,
                text_content=text_content,
            )

            return {
                "status": "success",
                "contract_id": target_contract_id,
                "message": "Versi dikonfirmasi dan diproses di latar belakang.",
            }

        # Legacy manual-link flow kept for older clients.
        legacy_payload = ConfirmVersionLinkRequest(**payload)
        if not legacy_payload.new_contract_id or not legacy_payload.parent_contract_id:
            raise HTTPException(status_code=400, detail="Missing pending version or legacy contract link parameters.")

        parent_res = supabase.table("contracts") \
            .select("id, title, version_count, matter_id, tenant_id") \
            .eq("id", legacy_payload.parent_contract_id) \
            .limit(1) \
            .execute()
        if not parent_res.data:
            raise HTTPException(status_code=404, detail="Parent contract not found.")

        parent = parent_res.data[0]
        tenant_id = parent.get("tenant_id") or base_tenant_id
        new_version_res = supabase.table("contract_versions") \
            .select("id") \
            .eq("contract_id", legacy_payload.new_contract_id) \
            .eq("tenant_id", tenant_id) \
            .gt("version_number", 0) \
            .order("version_number", desc=True) \
            .limit(1) \
            .execute()
        if not new_version_res.data:
            raise HTTPException(status_code=404, detail="New contract version not found.")

        next_version = _get_next_version_number(supabase, legacy_payload.parent_contract_id, tenant_id)
        version_row = new_version_res.data[0]
        supabase.table("contract_versions").update({
            "contract_id": legacy_payload.parent_contract_id,
            "version_number": next_version,
        }).eq("id", version_row["id"]).eq("tenant_id", tenant_id).execute()

        supabase.table("contracts").update({
            "version_count": next_version,
            "latest_version_id": version_row["id"],
        }).eq("id", legacy_payload.parent_contract_id).eq("tenant_id", tenant_id).execute()

        supabase.table("contracts").update({"status": "ARCHIVED"}) \
            .eq("id", legacy_payload.new_contract_id) \
            .eq("tenant_id", tenant_id) \
            .execute()

        return {
            "status": "success",
            "message": f"Contract linked as Version {next_version}.",
            "contract_id": legacy_payload.parent_contract_id,
            "new_version_number": next_version,
        }

    except HTTPException:
        raise
    except Exception as e:
        print(f"❌ Confirm Version Link Error: {e}")
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=str(e))


@router.post("/obligations/extract")
@limiter.limit("5/minute")
async def extract_obligations(
    request: Request,
    payload: ExtractObligationsRequest,
    claims: dict = Depends(verify_clerk_token),
    supabase: Client = Depends(get_tenant_supabase),
    qdrant_client: TenantQdrantClient = Depends(get_tenant_qdrant),
):
    try:
        tenant_id = claims["verified_tenant_id"]
        logger = TaskLogger(tenant_id=tenant_id, task_type="obligation_extraction", contract_id=payload.contract_id)

        # 1. Fetch Contract Text from Qdrant (NON-BLOCKING)
        contract_res, rules_res = await asyncio.gather(
            asyncio.to_thread(
                qdrant_client.scroll,
                collection_name=COLLECTION_NAME,
                scroll_filter=Filter(must=[
                    FieldCondition(key="contract_id", match=models.MatchValue(value=payload.contract_id)),
                ]),
                limit=100
            ),
            asyncio.to_thread(
                qdrant_client.scroll,
                collection_name="company_rules",
                limit=50
            ),
        )
        
        contract_text = ""
        for hit in contract_res[0]:
            contract_text += hit.payload.get("text", "") + "\n\n"
            
        if not contract_text.strip():
            raise HTTPException(status_code=404, detail="Contract text not found in vector database.")

        playbook_rules = ""
        if rules_res[0]:
            for hit in rules_res[0]:
                playbook_rules += f"- {hit.payload.get('rule_text', '')}\n"
        else:
            playbook_rules = "No custom playbook rules defined."

        # 3. Call OpenAI for Extraction & Compliance Check (NON-BLOCKING)
        system_prompt = """
You are an Elite Indonesian Corporate Lawyer AI. Read the provided CONTRACT TEXT.
Your task is to extract ONLY actionable obligations (things a party MUST DO, such as paying, delivering, or reporting). 
DO NOT extract general facts, background information, or declarations.

CRITICAL RULES:
1. ALWAYS write the output in INDONESIAN (Bahasa Indonesia).
2. Format each obligation clearly, starting with the responsible party. Example: "Vendor wajib..." or "Klien wajib...".
3. Evaluate each obligation against the provided COMPANY PLAYBOOK RULES.

Return a JSON object containing an array called 'obligations' with keys:
- 'description': (string) The actionable obligation in Indonesian.
- 'due_date': (string) The deadline, or 'N/A'.
- 'compliance_flag': (string) MUST BE 'SAFE' or 'CONFLICT' (if it violates the playbook).
"""
        user_prompt = f"COMPANY PLAYBOOK RULES:\n{playbook_rules}\n\nCONTRACT TEXT:\n{contract_text}"

        response = await async_chat_completion(
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt}
            ],
            response_format={"type": "json_object"}
        )
        
        raw_output = response.choices[0].message.content
        extracted_data = json.loads(raw_output)
        obligations_list = extracted_data.get("obligations", [])
        
        if not obligations_list:
            return {"status": "success", "message": "No obligations found.", "data": []}

        # 4. Save to Supabase
        insert_payload = []
        for ob in obligations_list:
            insert_payload.append({
                "tenant_id": tenant_id,
                "contract_id": payload.contract_id,
                "description": ob.get("description", "Unknown obligation"),
                "due_date": ob.get("due_date", None) if ob.get("due_date") != "N/A" else None,
                "status": "pending",
                "source": "AI",
                "compliance_flag": ob.get("compliance_flag", "SAFE")
            })
            
        # Each element in insert_payload already contains "tenant_id": tenant_id
        db_res = supabase.table("contract_obligations").insert([{**ob, "tenant_id": tenant_id} for ob in insert_payload]).execute()
        
        logger.complete(result_summary={"obligations_extracted": len(obligations_list)})
        return {
            "status": "success", 
            "message": f"Successfully extracted {len(obligations_list)} obligations.",
            "data": db_res.data
        }
    except Exception as e:
        if 'logger' in locals():
            logger.fail(e)
        print(f"Obligation Extraction Error: {e}")
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=str(e))

@router.patch("/{contract_id}/archive")
@limiter.limit("60/minute")
async def archive_contract(
    request: Request,
    contract_id: str,
    payload: ArchiveContractRequest,
    claims: dict = Depends(verify_clerk_token),
    supabase: Client = Depends(get_tenant_supabase)
):
    try:
        from datetime import datetime
        tenant_id = claims["verified_tenant_id"]
        
        update_payload = {
            "status": "ARCHIVED",
            "archive_reason": payload.archive_reason,
            "archived_at": datetime.utcnow().isoformat()
        }
        res = supabase.table("contracts").update(update_payload).eq("id", contract_id).eq("tenant_id", tenant_id).execute()
        
        if not res.data:
            raise HTTPException(status_code=404, detail="Contract not found or permission denied.")
            
        return {"status": "success", "message": "Contract archived successfully"}
    except HTTPException:
        raise
    except Exception as e:
        print(f"API Archive Contract Error: {e}")
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=str(e))

from typing import Optional

@router.get("/task-logs")
@limiter.limit("60/minute")
async def get_task_logs(
    request: Request,
    contract_id: Optional[str] = None,
    task_type: Optional[str] = None,
    status: Optional[str] = None,
    limit: int = 20,
    claims: dict = Depends(verify_clerk_token),
    supabase: Client = Depends(get_tenant_supabase),
):
    tenant_id = claims["verified_tenant_id"]
    
    query = supabase.table("task_execution_logs") \
        .select("*") \
        .eq("tenant_id", tenant_id) \
        .order("started_at", desc=True) \
        .limit(limit)
    
    if contract_id:
        query = query.eq("contract_id", contract_id)
    if task_type:
        query = query.eq("task_type", task_type)
    if status:
        query = query.eq("status", status)
    
    result = query.execute()
    return {"logs": result.data}

@router.get("/task-logs/{log_id}")
@limiter.limit("60/minute")
async def get_task_log_detail(
    request: Request,
    log_id: str,
    claims: dict = Depends(verify_clerk_token),
    supabase: Client = Depends(get_tenant_supabase),
):
    tenant_id = claims["verified_tenant_id"]
    
    result = supabase.table("task_execution_logs") \
        .select("*") \
        .eq("id", log_id) \
        .eq("tenant_id", tenant_id) \
        .single() \
        .execute()
    
    return {"log": result.data}
